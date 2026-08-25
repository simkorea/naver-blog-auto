# -*- coding: utf-8 -*-
"""
문서 텍스트 추출 — 계약서가 여기서 검색 가능해진다.

지원: PDF · 워드(docx) · 엑셀(xlsx) · 한글(hwpx/hwp) · 텍스트 · 이메일(eml)

PDF는 페이지 번호를 보존한다. "계약서 3쪽에 이 조항이 있다"고
가리킬 수 있어야 증거로 쓸 수 있기 때문이다.

라이브러리가 없으면 그 형식만 건너뛰고 나머지는 정상 처리한다.
"""
import re
import zipfile
from pathlib import Path

# 한 구간이 너무 길면 검색 결과에서 어디가 걸렸는지 보이지 않는다.
MAX_CHARS = 800


def _chunk(text: str, page_no=None, limit: int = MAX_CHARS) -> list[dict]:
    """긴 텍스트를 문단 경계에서 잘라 구간으로 만든다."""
    text = re.sub(r"[ \t]+", " ", text).strip()
    if not text:
        return []
    if len(text) <= limit:
        return [{"text": text, "page_no": page_no}]

    out, buf = [], ""
    for para in re.split(r"\n\s*\n|\n", text):
        para = para.strip()
        if not para:
            continue
        if len(buf) + len(para) + 1 > limit and buf:
            out.append({"text": buf.strip(), "page_no": page_no})
            buf = para
        else:
            buf = f"{buf}\n{para}" if buf else para
    if buf.strip():
        out.append({"text": buf.strip(), "page_no": page_no})
    return out


# ─────────────────────────────────────────────────────────
# 형식별 추출
# ─────────────────────────────────────────────────────────
# 이 글자 수보다 적게 나오면 스캔한 페이지로 본다.
# 스캔본은 글자가 그림 안에 있어 pdfplumber로는 아무것도 안 나온다.
SCANNED_THRESHOLD = 30


def from_pdf(path, ocr_scanned: bool = True) -> list[dict]:
    """
    페이지 번호를 보존해 추출. 표는 행 단위로 펼친다.

    계약서는 스캔본인 경우가 아주 흔하다. 그런 PDF는 글자가 그림으로
    들어 있어 텍스트 추출로는 한 글자도 안 나온다. 그대로 두면
    "계약서를 넣었는데 검색이 안 된다"가 된다.
    그래서 글자가 거의 없는 페이지는 이미지로 렌더링해 OCR로 읽는다.
    """
    try:
        import pdfplumber
    except ImportError:
        raise RuntimeError("PDF 추출에는 pdfplumber가 필요합니다 → pip install pdfplumber")

    segs = []
    scanned_pages = []

    with pdfplumber.open(path) as pdf:
        for page_no, page in enumerate(pdf.pages, 1):
            text = page.extract_text() or ""
            tables = page.extract_tables() or []

            if len(text.strip()) < SCANNED_THRESHOLD and not tables:
                scanned_pages.append(page_no)
                continue

            segs.extend(_chunk(text, page_no))

            # 계약서의 핵심 정보는 표 안에 있는 경우가 많다
            for table in tables:
                rows = [
                    " | ".join((c or "").strip() for c in row if c is not None)
                    for row in table
                ]
                body = "\n".join(r for r in rows if r.strip(" |"))
                if body.strip():
                    segs.extend(_chunk(f"[표]\n{body}", page_no))

    if scanned_pages and ocr_scanned:
        segs.extend(_ocr_pdf_pages(path, scanned_pages))

    return sorted(segs, key=lambda s: (s.get("page_no") or 0))


def _ocr_pdf_pages(path, page_numbers: list[int]) -> list[dict]:
    """
    스캔한 페이지를 이미지로 렌더링해 글자를 읽는다.

    렌더링 해상도를 200dpi로 올린다. 기본 72dpi로는 한글 인식률이
    크게 떨어진다 — 계약서 조항을 놓치면 의미가 없다.
    """
    try:
        import pypdfium2 as pdfium
    except ImportError:
        print("        [안내] 스캔 PDF 글자 읽기 건너뜀 → pip install pypdfium2")
        return []

    from .images import get_reader
    reader = get_reader()
    if reader is None:
        print("        [안내] 스캔 PDF 글자 읽기 건너뜀 (easyocr 미설치)")
        return []

    import tempfile

    out = []
    doc = pdfium.PdfDocument(str(path))
    try:
        for page_no in page_numbers:
            try:
                page = doc[page_no - 1]
            except IndexError:
                continue
            # scale 2.78 ≒ 200dpi (기본 72dpi 기준)
            bitmap = page.render(scale=200 / 72)
            image = bitmap.to_pil()

            with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp:
                tmp_path = tmp.name
            try:
                image.save(tmp_path)
                lines = reader.readtext(tmp_path, detail=1, paragraph=False)
                text = "\n".join(
                    item[1].strip() for item in lines
                    if len(item) > 1 and str(item[1]).strip()
                )
                if not text.strip():
                    continue
                confs = [float(i[2]) for i in lines if len(i) > 2]
                avg = sum(confs) / len(confs) if confs else 0.0
                for seg in _chunk(f"[스캔 페이지 · 글자 인식]\n{text}", page_no):
                    seg["confidence"] = round(avg, 3)
                    # 인식이 흐릿하면 원본 PDF를 직접 봐야 한다
                    seg["hallucination_risk"] = 1 if avg < 0.6 else 0
                    out.append(seg)
            finally:
                try:
                    Path(tmp_path).unlink()
                except OSError:
                    pass
    finally:
        doc.close()

    if out:
        print(f"        스캔 페이지 {len(page_numbers)}쪽을 글자 인식으로 읽었습니다")
    return out


def from_docx(path) -> list[dict]:
    try:
        import docx
    except ImportError:
        raise RuntimeError("워드 추출에는 python-docx가 필요합니다 → pip install python-docx")

    d = docx.Document(str(path))
    parts = [p.text for p in d.paragraphs if p.text.strip()]
    for table in d.tables:
        for row in table.rows:
            line = " | ".join(c.text.strip() for c in row.cells)
            if line.strip(" |"):
                parts.append(line)
    return _chunk("\n".join(parts))


def from_xlsx(path) -> list[dict]:
    try:
        import openpyxl
    except ImportError:
        raise RuntimeError("엑셀 추출에는 openpyxl이 필요합니다 → pip install openpyxl")

    wb = openpyxl.load_workbook(str(path), data_only=True, read_only=True)
    segs = []
    for ws in wb.worksheets:
        lines = []
        for row in ws.iter_rows(values_only=True):
            cells = [str(c).strip() for c in row if c is not None and str(c).strip()]
            if cells:
                lines.append(" | ".join(cells))
        if lines:
            segs.extend(_chunk(f"[시트: {ws.title}]\n" + "\n".join(lines)))
    wb.close()
    return segs


def from_hwpx(path) -> list[dict]:
    """
    한글 hwpx는 zip + xml 구조라 표준 라이브러리만으로 읽을 수 있다.
    (구형 .hwp 바이너리는 별도 처리)
    """
    texts = []
    with zipfile.ZipFile(path) as z:
        names = [n for n in z.namelist()
                 if n.startswith("Contents/") and n.endswith(".xml")]
        for name in sorted(names):
            raw = z.read(name).decode("utf-8", errors="ignore")
            # <hp:t>본문</hp:t> 형태로 텍스트가 들어 있다
            found = re.findall(r"<hp:t[^>]*>(.*?)</hp:t>", raw, re.S)
            if not found:
                found = re.findall(r"<[^>]*:t>(.*?)</[^>]*:t>", raw, re.S)
            for frag in found:
                frag = re.sub(r"<[^>]+>", "", frag)
                frag = (frag.replace("&lt;", "<").replace("&gt;", ">")
                            .replace("&amp;", "&").replace("&quot;", '"'))
                if frag.strip():
                    texts.append(frag.strip())
    return _chunk("\n".join(texts))


def from_hwp(path) -> list[dict]:
    """구형 한글 바이너리. pyhwp가 있으면 쓰고, 없으면 안내만 남긴다."""
    try:
        from hwp5.dataio import ParseError          # noqa: F401
        from hwp5.xmlmodel import Hwp5File
    except ImportError:
        raise RuntimeError(
            "구형 한글(.hwp)은 pyhwp가 필요합니다 → pip install pyhwp\n"
            "  설치가 어려우면 한글에서 파일을 열어 '다른 이름으로 저장 → hwpx' 또는 "
            "'PDF로 저장' 후 다시 등록하세요."
        )
    texts = []
    hwp = Hwp5File(str(path))
    try:
        for section in hwp.bodytext:
            for para in section.models():
                t = getattr(para, "content", None)
                if isinstance(t, str) and t.strip():
                    texts.append(t.strip())
    finally:
        hwp.close()
    return _chunk("\n".join(texts))


def from_txt(path) -> list[dict]:
    p = Path(path)
    for enc in ("utf-8-sig", "utf-8", "cp949"):
        try:
            return _chunk(p.read_text(encoding=enc))
        except (UnicodeDecodeError, LookupError):
            continue
    return _chunk(p.read_text(encoding="utf-8", errors="ignore"))


def from_eml(path) -> list[dict]:
    """이메일. 날짜·발신자가 확실해 타임라인의 좋은 뼈대가 된다."""
    import email
    from email import policy
    from email.utils import parsedate_to_datetime

    msg = email.message_from_bytes(Path(path).read_bytes(), policy=policy.default)
    header = "\n".join(filter(None, [
        f"보낸사람: {msg.get('From', '')}",
        f"받는사람: {msg.get('To', '')}",
        f"제목: {msg.get('Subject', '')}",
        f"날짜: {msg.get('Date', '')}",
    ]))

    body = ""
    if msg.is_multipart():
        for part in msg.walk():
            if part.get_content_type() == "text/plain":
                body += part.get_content()
    else:
        try:
            body = msg.get_content()
        except Exception:
            body = str(msg.get_payload())

    occurred = None
    try:
        occurred = parsedate_to_datetime(msg.get("Date")).isoformat(timespec="seconds")
    except Exception:
        pass

    segs = _chunk(f"{header}\n\n{body}")
    for s in segs:
        s["occurred_at"] = occurred
    return segs


_HANDLERS = {
    ".pdf": from_pdf, ".docx": from_docx, ".xlsx": from_xlsx, ".xls": from_xlsx,
    ".hwpx": from_hwpx, ".hwp": from_hwp, ".txt": from_txt, ".md": from_txt,
    ".eml": from_eml,
}


def extract(conn, source_row) -> int:
    """원본 하나를 추출해 DB에 넣는다."""
    from .. import db, integrity

    path = Path(source_row["path"])
    handler = _HANDLERS.get(path.suffix.lower())
    if handler is None:
        db.set_status(conn, source_row["id"], "failed",
                      f"지원하지 않는 형식: {path.suffix}")
        return 0

    try:
        segs = handler(path)
    except BaseException as e:
        db.set_status(conn, source_row["id"], "failed", str(e))
        integrity.log("extract_failed", source_id=source_row["id"], error=str(e))
        return 0

    for i, s in enumerate(segs, 1):
        s["seq"] = i
        s.setdefault("confidence", 1.0)      # 문자 기록은 전사 오류가 없다

    db.clear_segments(conn, source_row["id"])
    n = db.add_segments(conn, source_row["id"], segs)

    # 이메일이면 실제 발신 일시로 원본 시각을 정정
    dated = [s.get("occurred_at") for s in segs if s.get("occurred_at")]
    if dated:
        db.write(conn,
                 "UPDATE sources SET occurred_at = ?, occurred_at_est = 0 WHERE id = ?",
                 (min(dated), source_row["id"]))

    pages = {s.get("page_no") for s in segs if s.get("page_no")}
    detail = f"{n}개 구간" + (f" · {len(pages)}쪽" if pages else "")
    db.set_status(conn, source_row["id"], "extracted", detail)
    integrity.log("extract_document", source_id=source_row["id"],
                  segments=n, format=path.suffix)
    return n
