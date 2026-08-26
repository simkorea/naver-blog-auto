# -*- coding: utf-8 -*-
"""
법률검토메모 (docx) — 변호사에게 그대로 건네는 문서.

구성
  쟁점 → 내 주장 → 뒷받침 증거(파일·타임코드) → 근거 조문 원문 →
  관련 판례 → 예상 반론(불리한 발언)

여기 실리는 조문·판례는 전부 실존 검증을 통과한 것만이다.
검증에 실패한 것은 문서에 아예 들어가지 않고, 대신 몇 건이 걸러졌는지만
말미에 밝힌다. 무엇이 빠졌는지 숨기지 않기 위해서다.

이 문서는 내부 검토용이다. 수사기관 제출 패키지에는 기본으로 넣지 않는다.
"""
from datetime import datetime
from pathlib import Path

from ..law.commenter import DISCLAIMER
from ..search.hybrid import timecode

FONT = "맑은 고딕"


def _setup(doc, title, subtitle=""):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(10)
    style.element.rPr.rFonts.set(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia", FONT)

    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = h.add_run(title)
    r.font.size = Pt(18)
    r.font.bold = True

    if subtitle:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rr = p.add_run(subtitle)
        rr.font.size = Pt(10)
        rr.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rr = p.add_run(f"작성 {datetime.now():%Y년 %m월 %d일}")
    rr.font.size = Pt(9)
    rr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)


def _para(doc, text, size=10, bold=False, color=None, indent=0, italic=False):
    from docx.shared import Pt, RGBColor
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Pt(indent)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    if color:
        r.font.color.rgb = RGBColor(*color)
    return p


def write(conn, out_path, case_name: str = "", include_blocked_note: bool = True) -> Path:
    import docx
    from docx.shared import Pt

    from ..law import commenter

    doc = docx.Document()
    _setup(doc, "법률 검토 메모",
           case_name or "변호사 상담용 정리 자료 (내부 검토용)")

    _para(doc, f"※ {DISCLAIMER}", size=9, color=(0xC0, 0x00, 0x00))
    _para(doc,
          "※ 아래 조문·판례는 국가법령정보센터에서 원문을 조회해 실존을 확인한 것만 "
          "수록했습니다. 확인되지 않은 인용은 자동으로 제외됩니다.",
          size=9, color=(0x66, 0x66, 0x66))
    doc.add_paragraph()

    rows = commenter.list_comments(conn, status="verified")
    if not rows:
        _para(doc, "검증을 통과한 법률 코멘트가 없습니다.", size=10)
        doc.save(out_path)
        return Path(out_path)

    # 쟁점별로 묶는다
    by_issue = {}
    for r in rows:
        by_issue.setdefault(r["issue_name"] or "기타", []).append(r)

    for idx, (issue, items) in enumerate(by_issue.items(), 1):
        _para(doc, f"{idx}. {issue}", size=13, bold=True)

        good = [i for i in items if i["stance"] == "유리"]
        bad = [i for i in items if i["stance"] == "불리"]
        other = [i for i in items if i["stance"] not in ("유리", "불리")]

        if good:
            _para(doc, "◆ 나에게 유리한 정황", size=10, bold=True,
                  color=(0x1F, 0x70, 0x2F), indent=12)
            for it in good:
                _evidence(doc, it)
        if bad:
            _para(doc, "◆ 예상되는 반론 — 상대가 들고 올 수 있는 발언",
                  size=10, bold=True, color=(0xC0, 0x00, 0x00), indent=12)
            for it in bad:
                _evidence(doc, it)
        if other:
            _para(doc, "◆ 관련 정황", size=10, bold=True, indent=12)
            for it in other:
                _evidence(doc, it)

        # 이 쟁점에 걸린 근거 원문 (중복 제거)
        seen, refs = set(), []
        for it in items:
            for c in it["citations"]:
                key = (c["ref_type"], c["ref_id"])
                if key not in seen:
                    seen.add(key)
                    refs.append(c)
        if refs:
            _para(doc, "◆ 관련 법령·판례 원문", size=10, bold=True, indent=12)
            for c in refs:
                _citation(doc, c)

        doc.add_paragraph()

    if include_blocked_note:
        blocked = commenter.list_comments(conn, status="blocked")
        if blocked:
            _para(doc,
                  f"※ 인용 실존이 확인되지 않아 이 문서에서 제외된 항목이 "
                  f"{len(blocked)}건 있습니다.",
                  size=9, color=(0x88, 0x88, 0x88))

    _page_numbers(doc)
    out_path = Path(out_path)
    doc.save(out_path)
    return out_path


def _evidence(doc, item):
    """뒷받침 증거 한 건 — 어느 파일 어느 위치인지 반드시 적는다."""
    loc = (timecode(item["start_sec"]) if item["start_sec"] is not None
           else (f"{item['page_no']}쪽" if item["page_no"]
                 else (item["at"] or "")[:16].replace("T", " ")))
    where = Path(item["path"]).name if item.get("path") else "출처 없음"
    head = f"· {where}　{loc}"
    if item.get("speaker_label"):
        head += f"　({item['speaker_label']})"
    _para(doc, head, size=9, color=(0x44, 0x44, 0x44), indent=24)
    _para(doc, f"「{(item['text'] or '')[:400]}」", size=10, indent=36)
    if item.get("reasoning"):
        _para(doc, item["reasoning"], size=9, color=(0x66, 0x66, 0x66),
              indent=36, italic=True)


def _citation(doc, c):
    _para(doc, f"▪ {c['label']}" + (f"　{c['extra']}" if c.get("extra") else ""),
          size=10, bold=True, indent=24)
    _para(doc, c["body"][:1500], size=9, indent=36)
    if c.get("url"):
        _para(doc, c["url"], size=8, color=(0x00, 0x00, 0xC0), indent=36)


def _page_numbers(doc):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Pt

    footer = doc.sections[0].footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.font.size = Pt(9)
    for kind in ("begin", "PAGE", "end"):
        if kind == "PAGE":
            el = run._r.makeelement(qn("w:instrText"), {})
            el.set(qn("xml:space"), "preserve")
            el.text = " PAGE "
        else:
            el = run._r.makeelement(qn("w:fldChar"), {})
            el.set(qn("w:fldCharType"), kind)
        run._r.append(el)
