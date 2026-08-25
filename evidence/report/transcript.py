# -*- coding: utf-8 -*-
"""
재생안내서 · 녹취록 발췌 (docx).

재생안내서
  원본을 그대로 제출할 때 함께 내는 문서.
  "어느 파일 몇 분 몇 초부터 들으면 됩니다"를 표로 정리한다.
  받는 사람은 원본 파일을 열어 그 시각으로 이동하면 끝이다.

녹취록 발췌
  선택한 구간을 화자·타임코드와 함께 옮겨 적은 문서.
  머리말에 이것이 AI 자동 전사본임을 반드시 밝힌다. 법원 제출용
  정식 녹취록은 속기사무소 작성본이어야 하며, 이 문서는 그 전 단계의
  참고자료다. 이 고지를 빼면 문서 자체가 신뢰를 잃는다.

python-docx로 만든다. 사용자 PC(윈도우)에서 돌아가야 하므로
파이썬 밖의 의존성을 늘리지 않는다.
"""
from datetime import datetime
from pathlib import Path

FONT = "맑은 고딕"

DISCLAIMER = (
    "본 문서는 음성 인식 프로그램이 자동으로 옮겨 적은 전사본으로, 참고용입니다. "
    "법원 제출용 정식 녹취록은 공신력 있는 속기사무소에서 작성한 것이어야 합니다. "
    "'확인' 표시가 없는 구간은 원본 음성과 아직 대조하지 않은 부분이므로, "
    "인용 전 반드시 원본을 청취하여 확인하시기 바랍니다."
)

KIND_KR = {"audio": "녹음", "kakao": "카톡", "document": "문서",
           "image": "이미지", "email": "메일"}


def _setup(doc, title: str, subtitle: str = ""):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(10)
    # 한글 글꼴은 동아시아 폰트 속성까지 지정해야 실제로 적용된다
    style.element.rPr.rFonts.set(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia", FONT)

    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = h.add_run(title)
    run.font.size = Pt(18)
    run.font.bold = True

    if subtitle:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(subtitle)
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(f"작성 {datetime.now():%Y년 %m월 %d일}")
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)


def _notice(doc, text: str, color=(0xC0, 0x00, 0x00)):
    from docx.shared import Pt, RGBColor
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(*color)
    return p


def _page_numbers(doc):
    """바닥글에 페이지 번호. 여러 장짜리 문서에서는 필수다."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Pt

    footer = doc.sections[0].footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.font.size = Pt(9)
    for instr in ("begin", "PAGE", "end"):
        el = run._r.makeelement(qn("w:fldChar" if instr != "PAGE" else "w:instrText"), {})
        if instr == "PAGE":
            el.set(qn("xml:space"), "preserve")
            el.text = " PAGE "
        else:
            el.set(qn("w:fldCharType"), instr)
        run._r.append(el)


# ─────────────────────────────────────────────────────────
# 재생안내서
# ─────────────────────────────────────────────────────────
def write_guide(rows: list[dict], out_path, case_name: str = "") -> Path:
    """원본을 그대로 제출할 때 함께 내는 '어디를 들으세요' 안내서."""
    import docx
    from docx.shared import Pt, RGBColor

    doc = docx.Document()
    _setup(doc, "증거 재생 안내서",
           case_name or "원본 파일에서 아래 위치를 재생하시면 해당 내용을 확인하실 수 있습니다")

    _notice(doc,
            "※ 아래 '위치'는 첨부된 원본 파일 안에서의 재생 시각입니다. "
            "미디어 재생기에서 해당 시각으로 이동하시면 기재된 내용을 들으실 수 있습니다.",
            (0x33, 0x33, 0x33))
    _notice(doc, f"※ {DISCLAIMER}")
    doc.add_paragraph()

    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    heads = ["번호", "원본 파일", "위치", "화자", "내용", "확인"]
    widths = [0.6, 2.6, 1.5, 1.0, 5.4, 0.9]
    for i, (name, w) in enumerate(zip(heads, widths)):
        cell = table.rows[0].cells[i]
        cell.text = ""
        r = cell.paragraphs[0].add_run(name)
        r.font.bold = True
        r.font.size = Pt(9)

    from docx.shared import Inches
    for r in rows:
        cells = table.add_row().cells
        values = [str(r["no"]), r["file"], r["location"], r["speaker"],
                  r["text"], r["verified"]]
        for i, v in enumerate(values):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(v)
            run.font.size = Pt(9)
            if i == 2:                       # 위치는 굵게 — 이 문서의 핵심
                run.font.bold = True
            if i == 5 and v == "미검증":
                run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    for row in table.rows:
        for i, w in enumerate(widths):
            row.cells[i].width = Inches(w)

    doc.add_paragraph()
    _notice(doc,
            f"총 {len(rows)}건　·　원본 파일의 SHA-256 해시값은 별첨 '해시목록'을 참조하십시오.",
            (0x66, 0x66, 0x66))
    _page_numbers(doc)

    out_path = Path(out_path)
    doc.save(out_path)
    return out_path


# ─────────────────────────────────────────────────────────
# 녹취록 발췌
# ─────────────────────────────────────────────────────────
def write_transcript(conn, rows: list[dict], out_path,
                     context_lines: int = 2, case_name: str = "") -> Path:
    """
    선택 구간을 녹취록 형식으로 옮겨 적는다.

    앞뒤 맥락을 함께 싣는다. 한 문장만 떼어내면 의미가 뒤집힐 수 있고,
    "맥락을 잘라냈다"는 반박의 빌미가 되기 때문이다.
    """
    import docx
    from docx.shared import Pt, RGBColor

    from ..search.hybrid import context, timecode

    doc = docx.Document()
    _setup(doc, "녹취록 발췌", case_name)
    _notice(doc, f"※ {DISCLAIMER}")
    doc.add_paragraph()

    for r in rows:
        # ── 구간 머리 ────────────────────────
        p = doc.add_paragraph()
        run = p.add_run(f"[증거 {r['no']}]  {r['file']}　{r['location']}")
        run.font.bold = True
        run.font.size = Pt(11)

        meta = []
        if r["when"]:
            meta.append(f"일시 {r['when']}")
        if r["issue"]:
            meta.append(f"쟁점 {r['issue']}")
        if r["reason"]:
            meta.append(f"제출 사유 {r['reason']}")
        if meta:
            mp = doc.add_paragraph()
            mr = mp.add_run("　".join(meta))
            mr.font.size = Pt(9)
            mr.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        # ── 본문 (앞뒤 맥락 포함) ──────────────
        if r["kind"] == "audio" and context_lines and r.get("segment_id"):
            lines = context(conn, r["segment_id"], context_lines, context_lines)
        else:
            lines = [{"id": r.get("segment_id"), "text": r["text"],
                      "speaker_label": r["speaker"], "start_sec": r.get("start_sec")}]

        for line in lines:
            is_main = line.get("id") == r.get("segment_id")
            lp = doc.add_paragraph()
            lp.paragraph_format.left_indent = Pt(18)

            tc = timecode(line.get("start_sec")) if line.get("start_sec") is not None else ""
            who = line.get("speaker_label") or line.get("speaker") or ""
            head = f"{tc}　{who}　" if tc or who else ""

            hr = lp.add_run(head)
            hr.font.size = Pt(9)
            hr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

            tr = lp.add_run(line.get("text", "").replace("\n", " "))
            tr.font.size = Pt(10)
            if is_main:
                tr.font.bold = True

        # ── 신뢰도 · 확인 상태 ─────────────────
        note = f"전사 신뢰도: {r['reliability']}　·　{r['verified']}"
        np_ = doc.add_paragraph()
        np_.paragraph_format.left_indent = Pt(18)
        nr = np_.add_run(note)
        nr.font.size = Pt(8)
        nr.font.color.rgb = (RGBColor(0xC0, 0x00, 0x00)
                             if r["verified"] == "미검증" or "확인 필요" in r["reliability"]
                             else RGBColor(0x88, 0x88, 0x88))
        doc.add_paragraph()

    _page_numbers(doc)
    out_path = Path(out_path)
    doc.save(out_path)
    return out_path
