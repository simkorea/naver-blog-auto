# -*- coding: utf-8 -*-
"""
전체 내보내기 — 분석 결과 전부를 파일로 빼낸다.

프로그램 안에서만 볼 수 있으면 곤란하다. 변호사에게 통째로 넘기거나,
다른 프로그램에서 열어보거나, 인쇄해서 읽어야 할 때가 있다.

  · 전체 전사본 (docx/txt)  — 녹음 하나를 처음부터 끝까지
  · 전체 표 (xlsx/csv)      — 모든 구간을 표로
  · 타임라인 (html)         — 브라우저로 보는 시간순 정리
"""
import csv
from datetime import datetime
from pathlib import Path

from ..search.hybrid import timecode

DISCLAIMER = (
    "본 문서는 음성 인식 프로그램이 자동으로 옮겨 적은 전사본으로 참고용입니다. "
    "'미검증' 표시가 있는 구간은 원본 음성과 아직 대조하지 않은 부분입니다."
)


# ─────────────────────────────────────────────────────────
# 전체 전사본
# ─────────────────────────────────────────────────────────
def full_transcript_text(conn, source_id: int) -> str:
    src = conn.execute("SELECT * FROM sources WHERE id = ?", (source_id,)).fetchone()
    if not src:
        raise ValueError(f"자료를 찾을 수 없습니다: {source_id}")

    rows = conn.execute(
        """SELECT s.*, n.verified_by_ear, n.corrected_text
           FROM segments s LEFT JOIN notes n ON n.segment_id = s.id
           WHERE s.source_id = ? ORDER BY s.seq""", (source_id,)).fetchall()

    name = Path(src["path"]).name
    out = [
        "=" * 74,
        f"  전사본 — {name}",
        "=" * 74,
        "",
        f"  일시   : {(src['occurred_at'] or '')[:16].replace('T', ' ')}"
        + ("  (추정)" if src["occurred_at_est"] else ""),
        f"  상대방 : {src['counterparty'] or '-'}",
        f"  길이   : {timecode(src['duration_sec']) if src['duration_sec'] else '-'}",
        f"  구간   : {len(rows)}개",
        f"  모델   : {src['model_used'] or '-'}",
        f"  해시   : {src['sha256']}",
        "",
        f"  ※ {DISCLAIMER}",
        "",
        "-" * 74,
        "",
    ]

    for r in rows:
        text = r["corrected_text"] or r["text"]
        head_parts = []
        if r["start_sec"] is not None:
            head_parts.append(timecode(r["start_sec"]))
        elif r["page_no"]:
            head_parts.append(f"{r['page_no']}쪽")
        elif r["occurred_at"]:
            head_parts.append(r["occurred_at"][:16].replace("T", " "))
        if r["speaker_label"] or r["speaker"]:
            head_parts.append(r["speaker_label"] or r["speaker"])

        flags = []
        if r["corrected_text"]:
            flags.append("청취 후 수정")
        elif r["verified_by_ear"]:
            flags.append("청취 확인")
        elif src["kind"] == "audio":
            flags.append("미검증")
        if r["alt_mismatch"]:
            flags.append(r["mismatch_kind"] or "전사 불일치")
        if r["speaker_uncertain"]:
            flags.append("화자 불확실")

        head = "  ".join(head_parts)
        tail = f"   [{' · '.join(flags)}]" if flags else ""
        out.append(f"{head}{tail}")
        out.append(f"    {text}")
        if r["alt_mismatch"] and r["alt_text"]:
            out.append(f"    (2차 모델: {r['alt_text']})")
        out.append("")

    out.append("=" * 74)
    return "\n".join(out)


def full_transcript_docx(conn, source_id: int, out_path) -> Path:
    """전사본을 워드로. 인쇄해서 읽거나 변호사에게 넘길 때."""
    import docx
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    src = conn.execute("SELECT * FROM sources WHERE id = ?", (source_id,)).fetchone()
    rows = conn.execute(
        """SELECT s.*, n.verified_by_ear, n.corrected_text
           FROM segments s LEFT JOIN notes n ON n.segment_id = s.id
           WHERE s.source_id = ? ORDER BY s.seq""", (source_id,)).fetchall()

    doc = docx.Document()
    style = doc.styles["Normal"]
    style.font.name = "맑은 고딕"
    style.font.size = Pt(10)
    style.element.rPr.rFonts.set(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia",
        "맑은 고딕")

    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = h.add_run(f"전사본 — {Path(src['path']).name}")
    r.font.size = Pt(16)
    r.font.bold = True

    meta = doc.add_paragraph()
    mr = meta.add_run(
        f"일시 {(src['occurred_at'] or '')[:16].replace('T', ' ')}　"
        f"상대방 {src['counterparty'] or '-'}　구간 {len(rows)}개")
    mr.font.size = Pt(9)
    mr.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    dp = doc.add_paragraph()
    dr = dp.add_run(f"※ {DISCLAIMER}")
    dr.font.size = Pt(9)
    dr.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    doc.add_paragraph()

    for row in rows:
        p = doc.add_paragraph()
        head = []
        if row["start_sec"] is not None:
            head.append(timecode(row["start_sec"]))
        elif row["page_no"]:
            head.append(f"{row['page_no']}쪽")
        who = row["speaker_label"] or row["speaker"] or ""
        if who:
            head.append(who)

        hr = p.add_run("　".join(head) + "　")
        hr.font.size = Pt(9)
        hr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        hr.font.bold = bool(who)

        tr = p.add_run(row["corrected_text"] or row["text"])
        tr.font.size = Pt(10)

        flags = []
        if row["alt_mismatch"]:
            flags.append(row["mismatch_kind"] or "전사 불일치")
        if src["kind"] == "audio" and not row["verified_by_ear"] and not row["corrected_text"]:
            flags.append("미검증")
        if flags:
            fr = p.add_run(f"　[{' · '.join(flags)}]")
            fr.font.size = Pt(8)
            fr.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    out_path = Path(out_path)
    doc.save(out_path)
    return out_path


# ─────────────────────────────────────────────────────────
# 전체 표
# ─────────────────────────────────────────────────────────
COLUMNS = [
    ("자료", "file"), ("종류", "kind"), ("일시", "when"), ("위치", "location"),
    ("화자", "speaker"), ("내용", "text"), ("쟁점", "issue"), ("유불리", "stance"),
    ("신뢰도", "reliability"), ("확인", "verified"),
]


def all_rows(conn) -> list[dict]:
    """모든 구간을 표 한 줄씩으로."""
    from .locator import _build
    rows = conn.execute(
        """SELECT s.*, src.path, src.kind, src.sha256, src.counterparty,
                  src.occurred_at AS src_occurred_at, src.is_my_conversation,
                  n.verified_by_ear, n.corrected_text, n.verdict,
                  b.clip_start_sec, b.clip_end_sec, b.reason
           FROM segments s
           JOIN sources src ON src.id = s.source_id
           LEFT JOIN notes n ON n.segment_id = s.id
           LEFT JOIN basket b ON b.segment_id = s.id
           ORDER BY COALESCE(src.occurred_at, ''), src.id, s.seq"""
    ).fetchall()
    return _build(conn, [dict(r) for r in rows])


def all_csv(conn, out_path) -> Path:
    """엑셀에서 바로 열리게 UTF-8 BOM으로 저장한다."""
    rows = all_rows(conn)
    out_path = Path(out_path)
    with out_path.open("w", encoding="utf-8-sig", newline="") as f:
        w = csv.writer(f)
        w.writerow([label for label, _ in COLUMNS])
        for r in rows:
            w.writerow([r.get(key, "") for _, key in COLUMNS])
    return out_path


def all_xlsx(conn, out_path) -> Path:
    from .excel import write
    return write(all_rows(conn), out_path, title="전체 구간 목록")


# ─────────────────────────────────────────────────────────
# 타임라인 HTML
# ─────────────────────────────────────────────────────────
def timeline_html(conn, out_path, case_name: str = "") -> Path:
    from ..analyze import timeline

    events = timeline.events(conn, limit=5000)
    days = timeline.by_day(events)
    contras = timeline.contradictions(conn)

    def esc(t):
        return (str(t or "").replace("&", "&amp;").replace("<", "&lt;")
                .replace(">", "&gt;"))

    parts = [f"""<!doctype html><html lang="ko"><head><meta charset="utf-8">
<title>타임라인 — {esc(case_name) or '증거 정리'}</title><style>
:root{{color-scheme:light dark}}
body{{font-family:'맑은 고딕','Malgun Gothic',sans-serif;max-width:1000px;
margin:0 auto;padding:32px 20px;line-height:1.7;background:#fff;color:#1a1a1a}}
h1{{font-size:26px;margin:0 0 4px}}
.sub{{color:#777;font-size:13px;margin-bottom:28px}}
.day{{font-weight:700;font-size:17px;margin:32px 0 10px;padding-bottom:6px;
border-bottom:2px solid #1F3864}}
.ev{{margin:10px 0;padding:10px 14px;border-left:3px solid #ddd;background:#fafafa}}
.ev.good{{border-left-color:#1F702F;background:#f2f9f4}}
.ev.bad{{border-left-color:#C00000;background:#fdf3f3}}
.meta{{font-size:12px;color:#666}}
.meta b{{color:#1a1a1a}}
.flag{{color:#C00000;font-size:12px}}
.txt{{margin-top:4px}}
.contra{{border:2px solid #C00000;border-radius:8px;padding:16px;margin:16px 0;
background:#fff7f7}}
.contra h3{{margin:0 0 10px;font-size:16px;color:#C00000}}
.pair{{display:grid;grid-template-columns:1fr 1fr;gap:14px}}
.pair>div{{padding:10px;background:#fff;border-radius:6px;border:1px solid #eee}}
.note{{font-size:12px;color:#888;margin-top:40px;padding-top:14px;
border-top:1px solid #eee}}
@media(max-width:700px){{.pair{{grid-template-columns:1fr}}}}
</style></head><body>
<h1>{esc(case_name) or '증거 타임라인'}</h1>
<div class="sub">작성 {datetime.now():%Y년 %m월 %d일 %H:%M} · 사건 {len(events)}건</div>"""]

    if contras:
        parts.append("<h2 style='font-size:19px'>진술이 바뀐 정황</h2>")
        for c in contras[:10]:
            e, l = c["earlier"], c["later"]
            parts.append(f"""<div class="contra">
<h3>{esc(c['person'])} — {c['days']}일 간격으로 말이 달라졌습니다</h3>
<div class="pair">
<div><div class="meta"><b>{esc((e['at'] or '')[:10])}</b> · {esc(Path(e['path']).name)}
{' · ' + timecode(e['start_sec']) if e.get('start_sec') is not None else ''}</div>
<div class="txt">{esc(e['text'][:300])}</div></div>
<div><div class="meta"><b>{esc((l['at'] or '')[:10])}</b> · {esc(Path(l['path']).name)}
{' · ' + timecode(l['start_sec']) if l.get('start_sec') is not None else ''}</div>
<div class="txt">{esc(l['text'][:300])}</div></div>
</div></div>""")

    parts.append("<h2 style='font-size:19px;margin-top:36px'>시간순 정리</h2>")
    for day, evs in days:
        parts.append(f'<div class="day">{esc(day)}</div>')
        for ev in evs:
            cls = {"유리": "good", "불리": "bad"}.get(ev.get("stance"), "")
            loc = (timecode(ev["start_sec"]) if ev["start_sec"] is not None
                   else (f"{ev['page_no']}쪽" if ev["page_no"] else ""))
            who = ev["speaker_label"] or ev["speaker"] or ""
            flags = []
            if ev["occurred_at_est"]:
                flags.append("시각 추정")
            if ev["alt_mismatch"]:
                flags.append(ev["mismatch_kind"] or "전사 불일치")
            if ev["speaker_uncertain"]:
                flags.append("화자 불확실")

            parts.append(f"""<div class="ev {cls}">
<div class="meta">{esc(Path(ev['path']).name)}{' · ' + esc(loc) if loc else ''}
{' · <b>' + esc(who) + '</b>' if who else ''}
{' · ' + esc(ev['issue']) if ev['issue'] else ''}
{' <span class="flag">(' + esc(' · '.join(flags)) + ')</span>' if flags else ''}</div>
<div class="txt">{esc((ev['corrected_text'] or ev['text'])[:400])}</div></div>""")

    parts.append(f"""<div class="note">{esc(DISCLAIMER)}<br>
이 문서는 자료 정리용이며 법률 자문이 아닙니다.
증거 채택 여부와 제출 전략은 변호사와 상의하십시오.</div>
</body></html>""")

    out_path = Path(out_path)
    out_path.write_text("\n".join(parts), encoding="utf-8")
    return out_path
